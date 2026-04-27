"""
Generate RIS_Technical_Requirements.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"Q:\trm\Research\NGDS\viewer\RIS_Technical_Requirements.docx"

# ── Colors ──────────────────────────────────────────────────────────────────
BLUE_DARK  = "1F4E79"   # heading accent
BLUE_HDR   = "D5E8F0"   # table header fill
BLUE_ALT   = "F0F6FB"   # alternating row fill
WHITE      = "FFFFFF"

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Style helpers ────────────────────────────────────────────────────────────
def set_font(run, name="Arial", size=11, bold=False, color=None, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
         bold=False, size=11, color=None, italic=False, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def heading1(text):
    p = para(text, bold=True, size=14, color=BLUE_DARK, space_before=14, space_after=4)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BLUE_DARK)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    return para(text, bold=True, size=11, color=BLUE_DARK, space_before=8, space_after=2)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    r = p.add_run(text)
    set_font(r, size=10)
    return p

def numbered(text, num):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    r = p.add_run(f"{num}.\t{text}")
    set_font(r, size=10)
    return p

def body(text, space_after=6):
    return para(text, size=10, space_after=space_after)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)

def make_table(headers, rows, col_widths_in, alt=True):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Inches(col_widths_in[i])
        shade_cell(cell, BLUE_HDR)
        cell_text(cell, h, bold=True, size=10, color=BLUE_DARK)

    # Data rows
    for ri, row_data in enumerate(rows):
        tr = t.rows[ri + 1]
        fill = BLUE_ALT if (alt and ri % 2 == 1) else WHITE
        for ci, val in enumerate(row_data):
            cell = tr.cells[ci]
            cell.width = Inches(col_widths_in[ci])
            shade_cell(cell, fill)
            cell_text(cell, val, size=10)

    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    para()

para("Technical Requirements for", align=WD_ALIGN_PARAGRAPH.CENTER, size=16,
     bold=False, color="444444", space_after=2)
para("Multi-Site RT Plan Blinded Review Platform",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, color=BLUE_DARK, space_after=10)

# Decorative rule
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), BLUE_DARK)
pBdr.append(bot); pPr.append(pBdr)

para("Research Infrastructure Request — WashU RIS",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True, color="555555", space_after=24)

for _ in range(3):
    para()

for label, value in [
    ("Prepared by:", "Radiation Oncology Research, Washington University in St. Louis"),
    ("Date:", "April 2026"),
    ("Status:", "Draft for IT/RIS Review"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + "  ")
    set_font(r1, size=11, bold=True, color=BLUE_DARK)
    r2 = p.add_run(value)
    set_font(r2, size=11, color="333333")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1
# ════════════════════════════════════════════════════════════════════════════
heading1("1.  Project Overview")
body(
    "This document describes the infrastructure requirements for a web-based, "
    "multi-institutional radiation therapy (RT) plan review platform. The platform "
    "enables blinded physician ranking of competing RT plans across participating "
    "institutions as part of a prospective comparative planning study.", space_after=4
)
body(
    "The platform is pre-built and currently functional as a local prototype. "
    "This request covers the production hosting, storage, and networking requirements "
    "to deploy it at scale under WashU's research computing environment.", space_after=8
)
para("Key Characteristics", bold=True, size=10, space_after=2)
for b in [
    "Fully browser-based viewer — no software installation required by reviewers",
    "Pre-processed, de-identified DICOM data served as compressed binary files",
    "No patient health information is retained; all data is de-identified per HIPAA Safe Harbor",
    "Multi-institutional reviewer access via WashU SSO and InCommon federated identity",
    "Inbound data submission from external institutions via Globus file transfer",
    "Rankings and reviewer responses collected server-side and stored on RIS storage",
]:
    bullet(b)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2
# ════════════════════════════════════════════════════════════════════════════
heading1("2.  Storage Requirements")
make_table(
    ["Component", "Estimated Size", "Storage Tier", "Notes"],
    [
        ["Processed viewer files (CT + dose + structures, all plans)", "~50–200 MB per case", "Active (Tier 1)", "Served directly to browser; requires fast read access"],
        ["Raw DICOM archive (inbound from contributing sites)", "~500 MB–2 GB per case", "Archive (Tier 2)", "Cold storage; accessed only during pre-processing"],
        ["Rankings and response data", "<1 MB per case", "Active (Tier 1)", "JSON files; small but requires write access from web server"],
        ["Logs and audit trail", "~10 MB/month", "Active (Tier 1)", "Access logs for security and research audit purposes"],
    ],
    [2.6, 1.3, 1.1, 2.5]
)
body(
    "Initial study scope: approximately 100 cases (3 plans each). Total active storage estimate: "
    "20–50 GB. Total archive estimate: 50–200 GB. Storage should be expandable as the study "
    "scales to additional sites and cases.", space_after=4
)
body(
    "Request: A dedicated RIS storage allocation (project collection) with separate subdirectories "
    "for /intake, /processed, /rankings, and /archive, with independent access controls per subdirectory."
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3
# ════════════════════════════════════════════════════════════════════════════
heading1("3.  Compute Requirements")
body(
    "The platform requires a persistent, always-on web server process. This is distinct from batch "
    "compute jobs and represents the primary infrastructure question for the RIS team."
)

heading2("Option A — RIS Compute (Preferred)")
body(
    "A containerized web application (Node.js or Python/FastAPI, Docker or Singularity) running "
    "persistently on RIS compute infrastructure with direct filesystem access to RIS storage. "
    "This is the preferred option as it keeps all data within the RIS environment with no data "
    "movement to a separate host."
)
para("Requirements:", bold=True, size=10, space_before=4, space_after=2)
for b in [
    "1 persistent container or VM — 2–4 vCPU, 8 GB RAM",
    "Direct POSIX filesystem mount to RIS storage",
    "Outbound HTTPS capability",
    "Stable internal hostname for SSL certificate issuance",
]:
    bullet(b)

heading2("Option B — Separate IT-Hosted VM with RIS Storage Mount")
body(
    "If RIS Compute does not support persistent web services, a VM hosted by WashU IT with RIS "
    "storage mounted via NFS or similar. This is a common alternative pattern and remains fully "
    "compatible with the architecture."
)

heading2("Pre-Processing Pipeline")
body(
    "Separate from the web server, this pipeline processes incoming data. It can run as a "
    "scheduled batch job on RIS Compute:"
)
for b in [
    "Triggered on new data arrival in /intake/",
    "Runs DICOM validation, de-identification verification, and generates compressed viewer files",
    "Can run as a scheduled batch job (e.g., nightly) on RIS Compute",
    "Estimated runtime: 5–15 minutes per case",
    "Requirements: 4 vCPU, 16 GB RAM, access to both /intake/ and /processed/ storage",
]:
    bullet(b)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4
# ════════════════════════════════════════════════════════════════════════════
heading1("4.  Networking and Access Requirements")

heading2("4.1  Public HTTPS Access")
body(
    "The web application must be reachable over HTTPS from outside the WashU network, "
    "as reviewers at external institutions will access it from their own institutional networks."
)
for b in [
    "Public-facing IP or DNS entry (e.g., rtviewer.wustl.edu or similar)",
    "Valid SSL/TLS certificate (Let's Encrypt or WashU-issued)",
    "Coordination with WashU IT Networking for firewall rules and DNS",
]:
    bullet(b)

heading2("4.2  WashU Single Sign-On (SSO) Integration")
body("The application will authenticate users via WashU's SAML 2.0 identity provider.")
for b in [
    "Registration of the application as a SAML Service Provider with WashU Identity Services",
    "Exchange of SP metadata with identity.wustl.edu",
    "Required identity attributes: eduPersonPrincipalName (eppn), email, displayName, institutional affiliation",
    "Contact: WUIT-SA / WashU Identity and Access Management team",
]:
    bullet(b)

heading2("4.3  InCommon Federation (Multi-Site Expansion)")
body(
    "To allow reviewers from external institutions to authenticate with their home institution credentials:"
)
for b in [
    "Register the application as an InCommon Service Provider",
    "This enables any InCommon member institution (most major US research universities and AMCs) to federate without per-institution SSO configuration",
    "One-time registration process coordinated with WashU IT",
]:
    bullet(b)

heading2("4.4  Globus Collection for Data Intake")
body("For receiving de-identified DICOM data from contributing institutions:")
for b in [
    "Create a Globus collection mapped to the /intake/ subdirectory of the RIS storage allocation",
    "Configure write-only access for contributing institution Globus identities (managed per-site)",
    "Contributing sites use Globus Connect to transfer data — no special setup required on their end beyond standard Globus",
    "WashU RIS already operates a Globus endpoint; this request is for a project-specific collection within it",
]:
    bullet(b)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5
# ════════════════════════════════════════════════════════════════════════════
heading1("5.  Security and Access Control")
make_table(
    ["Role", "Authentication", "Data Access", "Notes"],
    [
        ["Administrator (study team)", "WashU SSO", "Full — all cases, rankings, randomization key", "PI and designated research staff only"],
        ["Reviewer (WashU)", "WashU SSO", "Assigned cases only, blinded plan labels", "Physicians participating as reviewers"],
        ["Reviewer (External)", "Home institution SSO via InCommon", "Assigned cases only, blinded plan labels", "Same access as WashU reviewers; institution verified via SSO assertion"],
        ["Submitter", "Globus identity (institutional)", "Write to /intake/ only; no read access to other cases", "Contributing sites; cannot review their own institution's cases"],
        ["Web server process", "Service account", "Read from /processed/; write to /rankings/ and /logs/", "Principle of least privilege"],
    ],
    [1.5, 1.6, 2.2, 2.2]
)
para("Additional Security Requirements", bold=True, size=10, space_after=2)
for b in [
    "All data in transit encrypted via TLS 1.2+",
    "All access events logged with timestamp, user identity, and resource accessed",
    "Audit logs retained for duration of study plus 3 years (research records requirement)",
    "No PHI stored at any layer; de-identification verified at intake before processing",
    "Reviewer whitelist managed by study administrator; access can be revoked individually",
]:
    bullet(b)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6
# ════════════════════════════════════════════════════════════════════════════
heading1("6.  Data Flow Summary")
for i, step in enumerate([
    "Contributing Site: De-identifies DICOM data locally, then transfers via Globus to /intake/{site}/{subject}/",
    "Intake Processing (RIS Compute, scheduled): Validates and processes DICOM, generates compressed viewer files, writes to /processed/{site}/{subject}/",
    "Web Server: Serves viewer files to authenticated browsers, receives ranking submissions, writes to /rankings/",
    "Study Team: Accesses aggregate rankings and audit logs via administrator interface",
], 1):
    numbered(step, i)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7
# ════════════════════════════════════════════════════════════════════════════
heading1("7.  Questions for the RIS Team")
questions = [
    "Can RIS Compute support a persistent (always-on) containerized web service, or is the infrastructure primarily designed for batch/scheduled jobs? If not, what is the recommended path for hosting a persistent web application with access to RIS storage?",
    "What is the process for requesting a named Globus collection mapped to a specific subdirectory of a RIS storage allocation, with per-user write permissions managed by the project team?",
    "Can the RIS environment support NFS or POSIX filesystem access from a WashU IT-hosted VM, as a fallback for Option B compute hosting?",
    "What is the estimated lead time for storage allocation provisioning once a project is approved?",
    "Is there a standard process for requesting a project-specific subdomain (e.g., rtviewer.research.wustl.edu) and SSL certificate for a research web application?",
    "What compliance documentation is required for hosting de-identified medical imaging research data on RIS? Is a data classification form or security review required?",
]
for i, q in enumerate(questions, 1):
    numbered(q, i)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8
# ════════════════════════════════════════════════════════════════════════════
heading1("8.  Summary of Requests")
make_table(
    ["Request", "Priority", "Dependency"],
    [
        ["RIS storage allocation (~50 GB active, ~200 GB archive, expandable)", "High", "Project approval"],
        ["Globus collection on RIS endpoint mapped to /intake/ with configurable write permissions", "High", "Storage allocation"],
        ["Persistent web application hosting (Option A: RIS Compute container; Option B: IT VM with RIS mount)", "High", "Networking/DNS"],
        ["Public HTTPS DNS entry and SSL certificate", "High", "Compute hosting decision"],
        ["SAML SP registration with WashU Identity Services", "High", "DNS/hosting in place"],
        ["InCommon SP federation registration", "Medium", "SAML integration working"],
        ["Scheduled batch compute for pre-processing pipeline", "Medium", "Storage allocation"],
    ],
    [4.5, 1.0, 2.0]
)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading1("Appendix: Technology Stack Summary")
make_table(
    ["Component", "Technology", "Notes"],
    [
        ["Web server", "Node.js (Express) or Python (FastAPI)", "Containerized via Docker/Singularity"],
        ["Frontend", "Vanilla HTML/CSS/JavaScript", "No framework dependencies; runs in any modern browser"],
        ["CT/Dose rendering", "HTML5 Canvas API", "No external rendering libraries required"],
        ["Data compression", "pako (gzip)", "Client-side decompression of binary volume data"],
        ["Authentication", "passport-saml (Node) or python3-saml", "SAML 2.0 SP implementation"],
        ["File transfer (inbound)", "Globus", "Industry standard for academic research data transfer"],
        ["Storage access", "POSIX filesystem (direct mount)", "No object storage API required"],
        ["Ranking persistence", "JSON files on RIS storage", "Simple, auditable, no database required for pilot scale"],
    ],
    [1.8, 2.2, 3.5]
)

doc.save(OUT)
print(f"Saved: {OUT}")
